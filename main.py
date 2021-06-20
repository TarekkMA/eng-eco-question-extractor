from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import csv


class Question:
    question: str
    answers: List[str]
    correct_index: int


class Chapter:
    title: str
    questions: List[Question]


def add_question(doc: Document, question: Question):
    doc.add_paragraph(question.question, style='List Number')
    for i, a in enumerate(question.answers):
        ap = doc.add_paragraph(style='List Bullet 2')
        ar = ap.add_run(a)
        if i == question.correct_index:
            ar.bold = True
            ar.underline = True


def read_chapter(number: int) -> Chapter:
    chapter = Chapter()
    chapter.title = f"الفصل {number}"
    chapter.questions = []
    with open(f'data/chapter{number}.csv') as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        for row in csv_reader:
            last_index = len(row) - 1
            for r in reversed(row):
                if row[last_index] != "":
                    break
                last_index -= 1
            question = Question()
            question.question = row[0]
            question.correct_index = int(row[last_index]) - 1
            question.answers = row[1:last_index]
            chapter.questions.append(question)

    return chapter


def main():
    document = Document()
    document.add_heading('بنك أسئلة - إقتصاد هندسي', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i in range(1, 6):
        chapter = read_chapter(i)
        document.add_heading(chapter.title, 1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        for q in chapter.questions:
            add_question(document, q)
        document.add_page_break()

    document.save('questions.docx')


if __name__ == "__main__":
    main()
